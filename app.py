# app.py

import os
import io
import re
import docx  # python-docx
import PyPDF2  # PyPDF2
import requests
from bs4 import BeautifulSoup
import google.generativeai as genai
from flask import Flask, request, render_template, jsonify, send_file
from dotenv import load_dotenv
from werkzeug.utils import secure_filename # For secure file handling
import subprocess
import tempfile
import uuid
import shutil
import time
from threading import Timer
import atexit

# --- Configuration ---
load_dotenv()  # Load environment variables from .env file

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 # Limit file size (e.g., 16MB)
# Only allow docx files
ALLOWED_EXTENSIONS = {'docx'}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# --- Gemini API Configuration ---
try:
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        print("Error: GEMINI_API_KEY not found in .env file.")
        gemini_model = None
    else:
        genai.configure(api_key=gemini_api_key)
        # Use a free, capable model like gemini-1.5-flash
        gemini_model = genai.GenerativeModel('gemini-1.5-flash')
        print("Gemini Model configured successfully.")
except Exception as e:
    print(f"Error configuring Gemini API: {e}")
    gemini_model = None

# --- Helper Functions ---

def allowed_file(filename):
    """Checks if the uploaded file extension is allowed."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def parse_resume(file_path):
    """
    Parses resume file (PDF or DOCX) and attempts to extract text and basic sections.
    Returns a dictionary of sections or {"ERROR": "message"}.
    """
    text = ""
    try:
        file_ext = file_path.rsplit('.', 1)[1].lower()
        if file_ext == 'docx':
            doc = docx.Document(file_path)
            full_text = [para.text for para in doc.paragraphs]
            text = '\n'.join(full_text)
        elif file_ext == 'pdf':
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    if reader.is_encrypted:
                         # Attempt to decrypt with an empty password, might fail
                         try:
                             reader.decrypt('')
                         except Exception as decrypt_err:
                             print(f"Could not decrypt PDF: {decrypt_err}")
                             return {"ERROR": "Could not decrypt password-protected PDF."}

                    full_text = []
                    for page_num, page in enumerate(reader.pages):
                        try:
                            full_text.append(page.extract_text())
                        except Exception as page_extract_err:
                             print(f"Warning: Could not extract text from PDF page {page_num + 1}: {page_extract_err}")
                             # Optionally add placeholder or skip page
                    text = '\n'.join(filter(None, full_text)) # Filter out None results if extraction failed on a page
            except PyPDF2.errors.PdfReadError as pdf_err:
                 print(f"Error reading PDF: {pdf_err}")
                 return {"ERROR": f"Invalid or corrupted PDF file: {pdf_err}"}
        else:
            return {"ERROR": "Unsupported file type"}

        if not text.strip():
            return {"ERROR": "Could not extract text from file. It might be image-based, empty, or corrupted."}

        # --- Basic Section Extraction (Improved Heuristic) ---
        parsed_data = {}
        # Regex to find potential section headers (e.g., all caps, or Title Case followed by newline)
        # This is still a heuristic and might misinterpret lines.
        # Prioritize common section names first.
        common_sections = [
            "SUMMARY", "PROFILE", "OBJECTIVE",
            "EXPERIENCE", "EMPLOYMENT HISTORY", "WORK HISTORY",
            "EDUCATION",
            "KEY SKILLS", "SKILLS", "TECHNICAL SKILLS", "COMPETENCIES",
            "PROJECTS",
            "CERTIFICATIONS", "LICENSES",
            "AWARDS", "HONORS",
            "PUBLICATIONS",
            "REFERENCES" # Often excluded or just a note
        ]
        # Normalize text slightly for matching
        normalized_text = "\n" + text.strip() + "\n" # Add newlines for boundary matching
        section_indices = {}

        # Find indices of common sections first
        for section in common_sections:
            # Search for the section name possibly followed by variations (e.g., space, colon) and newline
            # Case-insensitive search
            pattern = re.compile(r'\n\s*(' + re.escape(section) + r'[:\s]*)\n', re.IGNORECASE | re.MULTILINE)
            match = pattern.search(normalized_text)
            if match:
                # Store the start index and the matched header text (preserving original case if possible)
                section_indices[match.start(1)] = match.group(1).strip()

        # Sort found sections by their appearance order
        sorted_indices = sorted(section_indices.keys())

        # Extract content between sections
        last_index = 0
        current_section_name = "HEADER" # Content before the first recognized section

        for i, index in enumerate(sorted_indices):
            header_text = section_indices[index]
            content = normalized_text[last_index:index].strip()

            # Assign content to the previous section name
            if content:
                 # Normalize common section names for consistency
                 normalized_section_name = current_section_name.upper()
                 if "EXPERIENCE" in normalized_section_name or "EMPLOYMENT" in normalized_section_name or "WORK HISTORY" in normalized_section_name:
                     normalized_section_name = "EXPERIENCE"
                 elif "EDUCATION" in normalized_section_name:
                     normalized_section_name = "EDUCATION"
                 elif normalized_section_name == "KEY SKILLS":
                     normalized_section_name = "KEY SKILLS"  # Preserve KEY SKILLS as distinct
                 elif "SKILLS" in normalized_section_name or "TECHNICAL" in normalized_section_name or "COMPETENCIES" in normalized_section_name:
                     normalized_section_name = "SKILLS"
                 elif "SUMMARY" in normalized_section_name or "OBJECTIVE" in normalized_section_name or "PROFILE" in normalized_section_name:
                     normalized_section_name = "SUMMARY"
                 elif "PROJECTS" in normalized_section_name:
                     normalized_section_name = "PROJECTS"
                 # Add more normalizations if needed

                 parsed_data[normalized_section_name] = content

            # Update for the next iteration
            current_section_name = header_text # Use the found header as the next section name
            last_index = index + len(header_text) # Start next content search after the header

        # Add the content after the last found section
        final_content = normalized_text[last_index:].strip()
        if final_content:
             # Normalize the last section name as well
             normalized_section_name = current_section_name.upper()
             if "EXPERIENCE" in normalized_section_name or "EMPLOYMENT" in normalized_section_name or "WORK HISTORY" in normalized_section_name:
                 normalized_section_name = "EXPERIENCE"
             elif "EDUCATION" in normalized_section_name:
                 normalized_section_name = "EDUCATION"
             elif normalized_section_name == "KEY SKILLS":
                 normalized_section_name = "KEY SKILLS"  # Preserve KEY SKILLS as distinct
             elif "SKILLS" in normalized_section_name or "TECHNICAL" in normalized_section_name or "COMPETENCIES" in normalized_section_name:
                 normalized_section_name = "SKILLS"
             elif "SUMMARY" in normalized_section_name or "OBJECTIVE" in normalized_section_name or "PROFILE" in normalized_section_name:
                 normalized_section_name = "SUMMARY"
             elif "PROJECTS" in normalized_section_name:
                 normalized_section_name = "PROJECTS"

             parsed_data[normalized_section_name] = final_content

        # If no sections were found, put everything under "FULL_TEXT"
        if not parsed_data and text.strip():
             parsed_data["FULL_TEXT"] = text.strip()
             # Remove the default "HEADER" if it's empty and we have FULL_TEXT
             if "HEADER" in parsed_data and not parsed_data["HEADER"]:
                 del parsed_data["HEADER"]

        # Remove empty sections
        parsed_data = {k: v for k, v in parsed_data.items() if v and v.strip()}

        print(f"Parsed Sections: {list(parsed_data.keys())}")
        if not parsed_data:
             return {"ERROR": "Parsing finished, but no content sections were identified."}

        return parsed_data

    except Exception as e:
        print(f"Error parsing resume {file_path}: {e}")
        import traceback
        traceback.print_exc() # Print detailed traceback for debugging
        return {"ERROR": f"An unexpected error occurred during parsing: {e}"}

def escape_latex_text(text):
    """Basic LaTeX escaping for text content."""
    if not isinstance(text, str):
        text = str(text)
    
    # First, handle list formatting
    if '\\item' in text:
        # Ensure proper list environment
        if not (r'\begin{customitemize}' in text and r'\end{customitemize}' in text):
            text = r'\begin{customitemize}' + '\n' + text + '\n' + r'\end{customitemize}'
    
    # Then handle special characters
    # Order matters here! Escape backslash first.
    text = text.replace('\\', r'\textbackslash{}')
    text = text.replace('&', r'\&')
    text = text.replace('%', r'\%')
    text = text.replace('$', r'\$')
    text = text.replace('#', r'\#')
    text = text.replace('_', r'\_')
    text = text.replace('{', r'\{')
    text = text.replace('}', r'\}')
    text = text.replace('~', r'\textasciitilde{}')
    text = text.replace('^', r'\textasciicircum{}')
    
    # Handle common unicode bullets and numbers at the start of lines
    text = re.sub(r'^\s*([•●*–-])\s*(\d+\.)?\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*(\d+\.)\s*', '', text, flags=re.MULTILINE)
    
    # Fix any double-escaped backslashes
    text = text.replace(r'\\textbackslash{}', r'\textbackslash{}')
    text = text.replace(r'\\&', r'\&')
    text = text.replace(r'\\%', r'\%')
    text = text.replace(r'\\$', r'\$')
    text = text.replace(r'\\#', r'\#')
    text = text.replace(r'\\_', r'\_')
    text = text.replace(r'\\{', r'\{')
    text = text.replace(r'\\}', r'\}')
    
    return text

def convert_to_latex(parsed_data):
    """Converts parsed resume data into a basic LaTeX string using a template."""
    if "ERROR" in parsed_data:
        return parsed_data["ERROR"]
    if not parsed_data:
        return "ERROR: No parsed data provided for LaTeX conversion."

    # --- Minimal LaTeX Resume Template ---
    latex_string = r"""
\documentclass[11pt]{article}
\usepackage[margin=0.5in]{geometry}  % Reduced margins
\usepackage{hyperref}

% Remove page numbers
\pagenumbering{gobble}

% Basic formatting
\setlength{\parindent}{0pt}
\setlength{\parskip}{0.2em}  % Reduced paragraph spacing
\raggedright

% Section formatting
\renewcommand{\section}[1]{
  \vspace{0.2em}  % Reduced section spacing
  {\large\bfseries #1}
  \vspace{0.1em}  % Reduced spacing
  \hrule
  \vspace{0.2em}  % Reduced spacing
}

% List formatting
\renewcommand{\labelitemi}{$\bullet$}
\setlength{\itemsep}{0.1em}  % Reduced item spacing
\setlength{\parskip}{0.1em}  % Reduced paragraph spacing
\setlength{\topsep}{0.1em}  % Reduced top spacing for lists
\setlength{\partopsep}{0.1em}  % Reduced spacing around lists
\setlength{\parsep}{0.1em}  % Reduced spacing between paragraphs in lists

% Custom list environment to ensure consistent bullet points
\newenvironment{customitemize}
  {\begin{itemize}[label=$\bullet$]}
  {\end{itemize}}

\begin{document}

% --- Attempt to extract Name and Contact from HEADER ---
"""
    header_content = parsed_data.pop("HEADER", "") # Use and remove header data
    name = "Your Name Here" # Default
    contact_info = "" # Default

    if header_content:
        lines = header_content.split('\n')
        if lines:
            name = lines[0].strip() # Assume first line is name
            # Try to format remaining lines as contact info
            contact_items = []
            for line in lines[1:]:
                line = line.strip()
                if line:
                    # Basic check for email/phone/linkedin/github/portfolio markers
                    if '@' in line or 'mailto:' in line:
                        contact_items.append(f"Email: {escape_latex_text(line)}")
                    elif re.search(r'(\d{3}[-\.\s]??){2}\d{4}', line): # Basic phone number regex
                        contact_items.append(f"Phone: {escape_latex_text(line)}")
                    elif 'linkedin.com' in line:
                        # Extract username from LinkedIn URL
                        username = line.split('/')[-1]
                        contact_items.append(f"LinkedIn: \\href{{{line}}}{{{escape_latex_text(username)}}}")
                    elif 'github.com' in line:
                        # Extract username from GitHub URL
                        username = line.split('/')[-1]
                        contact_items.append(f"GitHub: \\href{{{line}}}{{{escape_latex_text(username)}}}")
                    elif 'http' in line: # Generic website/portfolio
                        contact_items.append(f"Website: \\href{{{line}}}{{{escape_latex_text(line)}}}")
                    else:
                        contact_items.append(escape_latex_text(line)) # Address or other info
            contact_info = " \\\\ ".join(contact_items) # Separate contact items with LaTeX newlines

    # Add Header block to LaTeX
    latex_string += f"""
\\begin{{center}}
    {{\\Large {escape_latex_text(name)}}} % Use extracted name
    \\vspace{{0.2em}} \\\\  % Reduced spacing
    {contact_info} % Add formatted contact info
\\end{{center}}
\\vspace{{0.3em}}  % Reduced spacing
"""

    # --- Add other sections ---
    # Define order (optional, but improves consistency)
    section_order = ["SUMMARY", "KEY SKILLS", "EXPERIENCE", "PROJECTS", "EDUCATION", "CERTIFICATIONS", "AWARDS", "PUBLICATIONS"]
    processed_sections = set()

    for section_name in section_order:
        if section_name in parsed_data:
            section_content = parsed_data[section_name]
            escaped_name = escape_latex_text(section_name.replace('_', ' ').title())
            escaped_content = escape_latex_text(section_content)

            latex_string += f"\n\\section{{{escaped_name}}}\n"

            # Special handling for Experience section
            if section_name == "EXPERIENCE":
                # Split content into individual experiences
                experiences = section_content.split('\n\n')
                for exp in experiences:
                    if exp.strip():
                        # Try to extract title and date
                        lines = exp.strip().split('\n')
                        if len(lines) >= 2:
                            # Remove any bullet points from the title
                            title = lines[0].strip().replace('•', '').replace('*', '').replace('-', '').strip()
                            date = lines[1].strip()
                            # Add experience heading with role name
                            latex_string += f"\\textbf{{{escape_latex_text(title)}}} \\hfill \\textit{{{escape_latex_text(date)}}}\n"
                            # Add remaining content as bullet points
                            if len(lines) > 2:
                                latex_string += r'\begin{customitemize}' + '\n'
                                for line in lines[2:]:
                                    if line.strip():
                                        latex_string += r'  \item ' + escape_latex_text(line.strip()) + '\n'
                                latex_string += r'\end{customitemize}' + '\n'
                        else:
                            # If format is unexpected, add as is
                            latex_string += f"{escaped_content}\n"
            elif section_name == "KEY SKILLS":
                # Special handling for KEY SKILLS - ensure each item is bulleted
                lines = [line.strip() for line in section_content.split('\n') if line.strip()]
                latex_string += r'\begin{customitemize}' + '\n'
                for line in lines:
                    # Remove any existing bullets and add LaTeX bullet
                    line = line.replace('•', '').replace('*', '').replace('-', '').strip()
                    latex_string += r'  \item ' + escape_latex_text(line) + '\n'
                latex_string += r'\end{customitemize}' + '\n'
            elif section_name == "EDUCATION":
                # Special handling for EDUCATION - format as paragraphs with dates
                lines = [line.strip() for line in section_content.split('\n') if line.strip()]
                for line in lines:
                    # Split the line into degree and date if possible
                    parts = line.split(' - ')
                    if len(parts) == 2:
                        degree = parts[0].strip()
                        date = parts[1].strip()
                        latex_string += f"\\textbf{{{escape_latex_text(degree)}}} \\hfill \\textit{{{escape_latex_text(date)}}}\n\n"
                    else:
                        latex_string += f"{escape_latex_text(line)}\n\n"
            else:
                # For other sections, use standard formatting
                lines = [line.strip() for line in section_content.split('\n') if line.strip()]
                is_likely_list = len(lines) > 1 and len(section_content) / len(lines) < 150

                if is_likely_list:
                    latex_string += r'\begin{customitemize}' + '\n'
                    for line in lines:
                        latex_string += r'  \item ' + escape_latex_text(line) + '\n'
                    latex_string += r'\end{customitemize}' + '\n'
                else:
                    latex_string += f"{escaped_content}\n"

            processed_sections.add(section_name)

    # Add any remaining sections not in the preferred order
    for section_name, section_content in parsed_data.items():
        if section_name not in processed_sections and section_name != "FULL_TEXT":
            escaped_name = escape_latex_text(section_name.replace('_', ' ').title())
            escaped_content = escape_latex_text(section_content)
            latex_string += f"\n\\section{{{escaped_name}}}\n{escaped_content}\n"

    latex_string += "\n\\end{document}\n"
    print("LaTeX conversion complete.")
    return latex_string


def scrape_job_description(url):
    """Scrapes the main job description text from a URL using basic heuristics."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Referer': 'https://www.google.com/' # Sometimes helps
        }
        response = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)

        # Check content type - only parse HTML
        content_type = response.headers.get('Content-Type', '').lower()
        if 'html' not in content_type:
             return f"ERROR: URL points to non-HTML content ({content_type})"

        soup = BeautifulSoup(response.content, 'html.parser')

        # --- Job Description Extraction Logic (Heuristics - Needs Improvement) ---
        job_text = ""
        selectors = [
            'div[class*="job-description"]', 'div[id*="job-description"]',
            'div[class*="jobdescription"]', 'div[id*="jobdescription"]',
            'div[class*="job-details"]', 'div[id*="job-details"]',
            'div[class*="jobDetails"]', 'div[id*="jobDetails"]',
            'section[class*="job-description"]', 'article[class*="job-description"]',
            'div[role="main"]', # Common on some platforms
            'main'
        ]

        potential_containers = []
        for selector in selectors:
            try:
                elements = soup.select(selector)
                if elements:
                    potential_containers.extend(elements)
            except Exception as e:
                print(f"Warning: CSS selector '{selector}' failed: {e}")


        if potential_containers:
            # Find the container with the most text content, preferring deeper elements
            best_container = max(potential_containers, key=lambda tag: len(tag.get_text(strip=True, separator=' ')))
            # Clean the chosen container before extracting text
            for element in best_container(['script', 'style', 'button', 'input', 'nav', 'header', 'footer', 'aside', 'form', 'figure', 'img']):
                element.decompose()
            job_text = best_container.get_text(separator='\n', strip=True)
        else:
            # Fallback: Get text from body, after cleaning common noise tags
            print("Warning: Could not find specific job description container via selectors. Falling back to cleaned body text.")
            body = soup.find('body')
            if body:
                for element in body(['script', 'style', 'nav', 'header', 'footer', 'aside', 'form', 'figure', 'img', 'button', 'input', 'svg']):
                    element.decompose()
                job_text = body.get_text(separator='\n', strip=True)
                # Basic cleaning (remove excessive blank lines)
                job_text = re.sub(r'\n\s*\n', '\n', job_text)
            else:
                return "ERROR: Could not find body tag in HTML."

        # Further cleaning: remove short lines that are likely remnants of UI elements
        job_text_lines = [line for line in job_text.split('\n') if len(line.strip()) > 10 or line.strip().endswith(':')]
        job_text = '\n'.join(job_text_lines)


        if len(job_text) < 150: # Increased threshold
             print(f"Warning: Extracted text seems too short ({len(job_text)} chars). Scraping might have failed or the description is minimal.")
             # Consider returning error if too short: return "ERROR: Extracted text too short, likely failed."

        print(f"Scraped job description (length: {len(job_text)} chars)")
        return job_text

    except requests.exceptions.Timeout:
        print(f"Error fetching URL {url}: Timeout")
        return "ERROR: The request timed out."
    except requests.exceptions.HTTPError as e:
        print(f"Error fetching URL {url}: HTTP {e.response.status_code}")
        return f"ERROR: Could not fetch URL (HTTP {e.response.status_code}). Check the URL or website permissions."
    except requests.exceptions.RequestException as e:
        print(f"Error fetching URL {url}: {e}")
        return f"ERROR: Could not fetch URL: {e}"
    except Exception as e:
        print(f"Error scraping job description: {e}")
        import traceback
        traceback.print_exc()
        return f"ERROR: Failed to scrape or parse the page: {e}"


def tailor_section_with_gemini(section_name, section_content, job_description):
    """Uses Gemini to tailor a resume section based on the job description."""
    if not gemini_model:
        return "ERROR: Gemini model not configured."
    if not section_content or not section_content.strip():
        return "ERROR: Section content is empty."
    if not job_description or not job_description.strip():
        return "ERROR: Job description is empty."

    print(f"Tailoring section '{section_name}' with Gemini...")

    # For KEY SKILLS and SUMMARY sections, return the original content with proper LaTeX formatting
    if section_name in ["KEY SKILLS", "SUMMARY"]:
        # Convert the content to a proper LaTeX list if it contains bullet points
        lines = [line.strip() for line in section_content.split('\n') if line.strip()]
        if section_name == "KEY SKILLS" or any(line.startswith(('•', '*', '-')) for line in lines):
            formatted_content = r'\begin{itemize}' + '\n'
            for line in lines:
                # Remove any existing bullets and add LaTeX bullet
                line = line.replace('•', '').replace('*', '').replace('-', '').strip()
                formatted_content += r'  \item ' + escape_latex_text(line) + '\n'
            formatted_content += r'\end{itemize}'
        else:
            # If no bullet points, keep as paragraph
            formatted_content = escape_latex_text(section_content)
        return formatted_content

    # Construct a more detailed prompt for other sections
    prompt = f"""
You are an expert resume writer and career coach. Your task is to rewrite the following resume section to be more impactful and specifically tailored to the provided job description.

**Instructions:**
1. **Analyze:** Carefully read the original resume section and the job description. Identify key skills, experiences, and keywords from the job description.
2. **Rewrite:** Revise the original section to highlight the candidate's experiences and skills that directly match the requirements and preferences mentioned in the job description.
3. **Keywords:** Naturally integrate relevant keywords from the job description into the rewritten text. Avoid keyword stuffing.
4. **Action Verbs:** Use strong action verbs to start bullet points or describe accomplishments.
5. **Format:** Output *only* the content of the section, without any section headers or titles. Use standard LaTeX formatting for lists (e.g., `\\item First point.\\n\\item Second point.`). Do not include explanations, apologies, or introductory phrases. The content will be automatically placed under the appropriate section header. If the original was a paragraph, keep it as a paragraph unless a list format is clearly better for the content and job description. If the original was a list, keep it as a list using `\\item`.
6. **Conciseness:** Keep the language clear, concise, and professional. Aim for 1-2 pages total resume length.
7. **No Suggestions:** Do not include any suggestions or placeholders. Only include actual content.
8. **No Headers:** Do not include any section headers or titles in your output. Only provide the content that should go under the section.
9. **Special Instructions for EXPERIENCE:** 
    - Preserve the exact job title and company name from the original
    - Keep the dates exactly as they appear in the original
    - Format each experience as:
      \\textbf{{Job Title}} \\hfill \\textit{{Date}}
      \\textbf{{Company Name}}
      \\begin{{itemize}}
      \\item First bullet point
      \\item Second bullet point
      \\end{{itemize}}
10. **LaTeX Escaping:** Make sure to escape special LaTeX characters like &, %, $, #, _, {{, }}, ~, ^ with a backslash. For example, write 'R\\&D' instead of 'R&D'.
11. **List Formatting:** For lists, use this exact format:
    \\begin{{itemize}}
    \\item First item
    \\item Second item
    \\end{{itemize}}

**Job Description:**
---
{job_description[:3000]}
---
(Job description truncated if too long)

**Original Resume Section ({section_name}):**
---
{section_content}
---

**Rewritten Resume Section Content (LaTeX format only, no headers):**
"""

    try:
        response = gemini_model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=1024,
                temperature=0.7
            )
        )

        if not response.candidates:
            feedback = response.prompt_feedback
            print(f"Gemini Warning: No candidates generated. Feedback: {feedback}")
            if feedback.block_reason == 'SAFETY':
                return f"ERROR: Content generation blocked due to safety concerns: {feedback.safety_ratings}"
            else:
                return f"ERROR: AI model did not generate a response. Reason: {feedback.block_reason or 'Unknown'}"

        if response.candidates[0].content.parts:
            tailored_content = response.text.strip()
            if not tailored_content:
                print("Gemini Warning: Generated content is empty.")
                return "ERROR: AI model returned empty content."

            # Remove any suggestion patterns
            tailored_content = re.sub(r'\[Suggest.*?\]', '', tailored_content)
            tailored_content = re.sub(r'\(e\.g\.,.*?\)', '', tailored_content)
            tailored_content = re.sub(r'\(add.*?\)', '', tailored_content)

            # Remove text in brackets, parentheses that suggest edits/additions
            tailored_content = re.sub(r'\[.*?\]', '', tailored_content) # Remove any [text]
            tailored_content = re.sub(r'\(.*?\)', '', tailored_content) # Remove any (text)
            tailored_content = re.sub(r'<.*?>', '', tailored_content)   # Remove any <text>
            # Remove common placeholder patterns
            tailored_content = re.sub(r'(?i)(insert|add|include|write|describe|specify|enter|input|your|paste|put)(\s+.*?)(here|below|above)', '', tailored_content)
            # Remove "TODO" style comments
            tailored_content = re.sub(r'(?i)(todo|note|fixme|xxx|placeholder).*?\n', '', tailored_content)
            # Remove lines with common placeholder indicators
            tailored_content = re.sub(r'.*\.\.\.$', '', tailored_content, flags=re.MULTILINE)
            tailored_content = re.sub(r'.*_+.*', '', tailored_content, flags=re.MULTILINE)
            # Clean up any resulting empty lines
            tailored_content = re.sub(r'\n\s*\n+', '\n\n', tailored_content)
            tailored_content = tailored_content.strip()

            # Remove any section headers that might have been included
            tailored_content = re.sub(r'^\\section.*?$', '', tailored_content, flags=re.MULTILINE)
            tailored_content = re.sub(r'^\\section\*.*?$', '', tailored_content, flags=re.MULTILINE)

            # Ensure proper list formatting
            if '\\item' in tailored_content:
                # If content contains \item but no list environment, wrap it in itemize
                if not (r'\begin{itemize}' in tailored_content and r'\end{itemize}' in tailored_content):
                    tailored_content = r'\begin{itemize}' + '\n' + tailored_content + '\n' + r'\end{itemize}'

            # Fix any double-escaped backslashes
            tailored_content = tailored_content.replace(r'\\textbackslash{}', r'\textbackslash{}')
            tailored_content = tailored_content.replace(r'\\&', r'\&')
            tailored_content = tailored_content.replace(r'\\%', r'\%')
            tailored_content = tailored_content.replace(r'\\$', r'\$')
            tailored_content = tailored_content.replace(r'\\#', r'\#')
            tailored_content = tailored_content.replace(r'\\_', r'\_')
            tailored_content = tailored_content.replace(r'\\{', r'\{')
            tailored_content = tailored_content.replace(r'\\}', r'\}')

            # Final validation - if content is empty after cleaning, return error
            if not tailored_content.strip():
                return "ERROR: Generated content was empty or contained only placeholders."

            print(f"Gemini tailoring successful for section '{section_name}'.")
            return tailored_content
        else:
            print(f"Gemini Warning: Response candidate has no content parts. Candidate: {response.candidates[0]}")
            return "ERROR: AI model response structure invalid (no content parts)."

    except Exception as e:
        print(f"Error calling Gemini API for section '{section_name}': {e}")
        import traceback
        traceback.print_exc()
        return f"ERROR: Failed to interact with AI model: {e}"


def update_latex(original_latex, section_name, tailored_content):
    """
    Updates a specific section in the LaTeX string with tailored content.
    Handles both \section and \section* commands.
    """
    # Prepare the section name as it appears in the LaTeX \section command
    latex_section_name = section_name.replace('_', ' ').title()
    
    # Try both \section and \section* patterns
    section_patterns = [
        f"\\\\section\\*{{{latex_section_name}}}",  # \section*{Section Name}
        f"\\\\section{{{latex_section_name}}}"      # \section{Section Name}
    ]
    
    # Combine patterns with OR operator
    pattern = '|'.join(section_patterns)
    
    # Use regex for more robust finding, ignoring whitespace variations around the marker
    # Pattern: marker, followed by optional whitespace, then capture the content until the next \section or \end{document}
    # DOTALL allows '.' to match newlines. Use non-greedy '.*?'
    full_pattern = re.compile(
        f"({pattern})\\s*(.*?)\\s*(?=\\\\section|\\\\end{{document}})",
        re.DOTALL | re.IGNORECASE  # Ignore case for section marker
    )

    match = full_pattern.search(original_latex)

    if not match:
        print(f"Warning: Could not find section marker for '{section_name}' in LaTeX. Skipping update.")
        return original_latex

    # Get the section marker and content
    section_marker = match.group(1)
    old_content = match.group(2)

    # Ensure tailored content has appropriate spacing
    formatted_tailored_content = "\n" + tailored_content.strip() + "\n"

    # Replace the old content with the new tailored content, keeping the section marker
    updated_latex = original_latex[:match.start(2)] + formatted_tailored_content + original_latex[match.end(2):]

    print(f"Successfully updated section: {section_name}")
    return updated_latex


def check_latex_packages():
    """Check if required LaTeX packages are installed and install them if needed."""
    required_packages = [
        'latex-base',
        'latex-fonts-recommended',
        'latex-fonts-extra',
        'texlive-latex-extra',
        'texlive-fonts-recommended',
        'texlive-fonts-extra',
        'texlive-latex-recommended'
    ]
    
    try:
        # Check if pdflatex is installed
        subprocess.run(['pdflatex', '--version'], capture_output=True, check=True)
        
        # Try to compile a minimal test document
        with tempfile.TemporaryDirectory() as temp_dir:
            test_tex = os.path.join(temp_dir, 'test.tex')
            with open(test_tex, 'w') as f:
                f.write(r"""
\documentclass{article}
\usepackage{geometry}
\usepackage{titlesec}
\usepackage{marvosym}
\usepackage{enumitem}
\usepackage{hyperref}
\begin{document}
Test
\end{document}
""")
            try:
                subprocess.run(['pdflatex', '-interaction=nonstopmode', test_tex], 
                             cwd=temp_dir, 
                             capture_output=True, 
                             check=True)
                print("LaTeX packages check passed")
                return True
            except subprocess.CalledProcessError as e:
                print(f"LaTeX test compilation failed: {e.stderr}")
                return False
    except subprocess.CalledProcessError:
        print("pdflatex not found")
        return False
    except Exception as e:
        print(f"Error checking LaTeX packages: {e}")
        return False

def cleanup_pdf_file(file_path, delay=300):  # 5 minutes delay
    """Delete a PDF file after a specified delay."""
    def delete_file():
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Cleaned up PDF file: {file_path}")
        except Exception as e:
            print(f"Error cleaning up PDF file {file_path}: {e}")
    
    Timer(delay, delete_file).start()

def check_libreoffice():
    """Check if LibreOffice is installed and return the path to soffice."""
    try:
        # Try common installation paths
        possible_paths = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
            '/usr/bin/soffice',  # Linux
            '/usr/lib/libreoffice/program/soffice',  # Linux alternative
            'C:\\Program Files\\LibreOffice\\program\\soffice.exe',  # Windows
            'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe'  # Windows 32-bit
        ]
        
        # First try which/where command
        try:
            if os.name == 'nt':  # Windows
                result = subprocess.run(['where', 'soffice'], capture_output=True, text=True, check=True)
            else:  # Unix-like
                result = subprocess.run(['which', 'soffice'], capture_output=True, text=True, check=True)
            return result.stdout.strip()
        except subprocess.CalledProcessError:
            pass
        
        # Then check common paths
        for path in possible_paths:
            if os.path.exists(path):
                return path
                
        return None
    except Exception as e:
        print(f"Error checking for LibreOffice: {e}")
        return None

def convert_to_pdf(docx_path, output_dir):
    """Convert DOCX to PDF using available tools."""
    soffice_path = check_libreoffice()
    
    if soffice_path:
        try:
            pdf_path = os.path.join(output_dir, 'updated.pdf')
            subprocess.run([
                soffice_path,
                '--headless',
                '--convert-to', 'pdf:writer_pdf_Export',
                '--outdir', output_dir,
                docx_path
            ], check=True)
            
            # Find the generated PDF
            if not os.path.exists(pdf_path):
                for f in os.listdir(output_dir):
                    if f.endswith('.pdf'):
                        pdf_path = os.path.join(output_dir, f)
                        break
            return pdf_path
        except Exception as e:
            print(f"LibreOffice conversion failed: {e}")
            return None
    
    # Fallback: Try using python-docx2pdf if available
    try:
        import docx2pdf
        pdf_path = os.path.join(output_dir, 'updated.pdf')
        docx2pdf.convert(docx_path, pdf_path)
        return pdf_path
    except ImportError:
        print("docx2pdf not installed")
    except Exception as e:
        print(f"docx2pdf conversion failed: {e}")
    
    return None

# --- Flask Routes ---

@app.route('/')
def index():
    """Serves the main HTML page."""
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_resume():
    """Handles file upload, parsing, and returns tailored summary and skills."""
    if 'resume' not in request.files:
        return jsonify({"error": "No resume file part in the request."}), 400

    file = request.files['resume']
    manual_jd = request.form.get('job_description', '').strip()

    if file.filename == '':
        return jsonify({"error": "No file selected."}), 400
    if not manual_jd:
        return jsonify({"error": "Job description is required."}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        if not filename:
            filename = "uploaded_resume.docx"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        try:
            file.save(file_path)
            print(f"File saved to: {file_path}")
            # Parse resume
            parsed_data = parse_resume(file_path)
            if "ERROR" in parsed_data:
                return jsonify({"error": f"Parsing failed: {parsed_data['ERROR']}"}), 400
            if not parsed_data:
                return jsonify({"error": "Parsing failed: No sections found in the resume."}), 400

            # Get both SUMMARY and SKILLS sections
            summary = parsed_data.get("SUMMARY")
            skills = parsed_data.get("SKILLS") or parsed_data.get("KEY SKILLS")  # Try both common names
            
            if not summary:
                return jsonify({"error": "No SUMMARY section found in the resume."}), 400
            if not skills:
                return jsonify({"error": "No SKILLS section found in the resume."}), 400

            # Tailor both sections using Gemini
            if not gemini_model:
                return jsonify({"error": "AI model is not configured."}), 500

            # Tailor summary
            summary_prompt = f"""
You are an expert resume writer. Rewrite the following resume summary so that it is highly tailored to the provided job description and optimized to pass Applicant Tracking Systems (ATS). Use keywords from the job description naturally. Do not include any section headers or explanations. Return only the rewritten summary text.

Job Description:
{manual_jd}

Original Summary:
{summary}

Rewritten Summary:
"""
            # Tailor skills
            skills_prompt = f"""
You are an expert resume writer. Rewrite the following skills section to be highly tailored to the provided job description. Focus on:
1. Prioritizing skills that match the job requirements
2. Using the exact terminology from the job description
3. Grouping related skills together
4. Removing irrelevant skills
5. Adding any missing critical skills from the job description that the candidate likely has

Format the output as a bullet-point list, with each skill on a new line starting with a bullet point (•). Do not include any section headers or explanations.

Job Description:
{manual_jd}

Original Skills:
{skills}

Rewritten Skills (bullet points only):
"""

            try:
                # Generate tailored summary
                summary_response = gemini_model.generate_content(
                    summary_prompt,
                    generation_config=genai.types.GenerationConfig(
                        max_output_tokens=512,
                        temperature=0.7
                    )
                )
                tailored_summary = summary_response.text.strip() if summary_response and hasattr(summary_response, 'text') else None

                # Generate tailored skills
                skills_response = gemini_model.generate_content(
                    skills_prompt,
                    generation_config=genai.types.GenerationConfig(
                        max_output_tokens=512,
                        temperature=0.7
                    )
                )
                tailored_skills = skills_response.text.strip() if skills_response and hasattr(skills_response, 'text') else None

                if not tailored_summary or not tailored_skills:
                    return jsonify({"error": "AI did not return complete content."}), 500

                return jsonify({
                    "tailored_summary": tailored_summary,
                    "tailored_skills": tailored_skills
                })

            except Exception as e:
                print(f"Error calling Gemini: {e}")
                return jsonify({"error": f"AI error: {e}"}), 500

        finally:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except OSError as e:
                    print(f"Error removing file {file_path}: {e}")
    else:
        return jsonify({"error": "Invalid file type. Only .docx is allowed."}), 400

@app.route('/preview', methods=['POST'])
def preview_latex():
    """Handles LaTeX preview and PDF generation."""
    print("Received preview request")
    latex_content = request.form.get('latex', '').strip()
    print(f"LaTeX content length: {len(latex_content)}")
    
    if not latex_content:
        print("Error: No LaTeX content provided")
        return jsonify({"error": "No LaTeX content provided"}), 400

    try:
        # Create a temporary directory for LaTeX compilation
        with tempfile.TemporaryDirectory() as temp_dir:
            print(f"Created temporary directory: {temp_dir}")
            # Generate unique filenames
            base_name = str(uuid.uuid4())
            tex_file = os.path.join(temp_dir, f"{base_name}.tex")
            pdf_file = os.path.join(temp_dir, f"{base_name}.pdf")
            log_file = os.path.join(temp_dir, f"{base_name}.log")

            # Write LaTeX content to file
            with open(tex_file, 'w', encoding='utf-8') as f:
                f.write(latex_content)
            print(f"Wrote LaTeX content to: {tex_file}")

            # Compile LaTeX to PDF using pdflatex
            try:
                print("Attempting to compile LaTeX...")
                result = subprocess.run(['pdflatex', '-interaction=nonstopmode', tex_file], 
                                     cwd=temp_dir, 
                                     capture_output=True, 
                                     text=True, 
                                     check=True)
                print("LaTeX compilation successful")
            except subprocess.CalledProcessError as e:
                print(f"LaTeX compilation failed: {e.stderr}")
                # Read the log file for more detailed error information
                error_details = ""
                if os.path.exists(log_file):
                    with open(log_file, 'r', encoding='utf-8') as f:
                        error_details = f.read()
                    print(f"LaTeX log file contents: {error_details}")
                return jsonify({
                    "error": f"LaTeX compilation failed: {e.stderr}",
                    "details": error_details
                }), 400

            # Check if PDF was generated
            if not os.path.exists(pdf_file):
                print("Error: PDF file was not generated")
                # Read the log file for error information
                error_details = ""
                if os.path.exists(log_file):
                    with open(log_file, 'r', encoding='utf-8') as f:
                        error_details = f.read()
                    print(f"LaTeX log file contents: {error_details}")
                return jsonify({
                    "error": "PDF generation failed",
                    "details": error_details
                }), 500

            # Create a static directory for PDFs if it doesn't exist
            static_dir = os.path.join(app.root_path, 'static', 'pdfs')
            os.makedirs(static_dir, exist_ok=True)
            print(f"Created static directory: {static_dir}")

            # Copy the PDF to the static directory
            static_pdf_path = os.path.join(static_dir, f"{base_name}.pdf")
            with open(pdf_file, 'rb') as src, open(static_pdf_path, 'wb') as dst:
                dst.write(src.read())
            print(f"Copied PDF to: {static_pdf_path}")

            # Schedule cleanup of the PDF file
            cleanup_pdf_file(static_pdf_path)

            # Generate URLs for preview and download
            preview_url = f"/static/pdfs/{base_name}.pdf"
            download_url = f"/download/{base_name}.pdf"
            print(f"Generated URLs - Preview: {preview_url}, Download: {download_url}")

            return jsonify({
                "preview_url": preview_url,
                "download_url": download_url
            })

    except Exception as e:
        print(f"Error in preview generation: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/download/<filename>')
def download_pdf(filename):
    """Handles PDF downloads."""
    try:
        pdf_path = os.path.join(app.root_path, 'static', 'pdfs', filename)
        if not os.path.exists(pdf_path):
            return jsonify({"error": "PDF not found"}), 404
        
        # Schedule cleanup after download
        cleanup_pdf_file(pdf_path)
        
        return send_file(pdf_path, as_attachment=True, download_name=filename)
    except Exception as e:
        print(f"Error in PDF download: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/download-docx', methods=['POST'])
def download_docx():
    """
    Creates a new document preserving all formatting from the original,
    replacing only the summary and skills sections with tailored content.
    """
    if 'resume' not in request.files or 'tailored_summary' not in request.form or 'tailored_skills' not in request.form:
        return jsonify({'error': 'Missing file or tailored sections.'}), 400
    resume_file = request.files['resume']
    tailored_summary = request.form['tailored_summary']
    tailored_skills = request.form['tailored_skills']
    if not resume_file.filename.lower().endswith('.docx'):
        return jsonify({'error': 'Only .docx files are supported.'}), 400
    try:
        # Save uploaded docx to temp file
        with tempfile.TemporaryDirectory() as tmpdir:
            orig_docx_path = os.path.join(tmpdir, 'orig.docx')
            resume_file.save(orig_docx_path)
            
            # Load the original document
            doc = docx.Document(orig_docx_path)
            
            def find_section_with_formatting(section_name):
                start_idx = None
                end_idx = len(doc.paragraphs)
                section_style = None
                content_styles = []
                
                # First pass: find section and collect styles
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip().upper().startswith(section_name):
                        start_idx = i
                        section_style = para.style
                        # Find where section ends (next all-caps heading or end)
                        for j in range(i+1, len(doc.paragraphs)):
                            t = doc.paragraphs[j].text.strip()
                            if t.isupper() and len(t) > 2 and not t.startswith(section_name):
                                end_idx = j
                                break
                        # Collect styles of content paragraphs
                        for j in range(i+1, end_idx):
                            if doc.paragraphs[j].text.strip():
                                content_styles.append(doc.paragraphs[j].style)
                        break
                return start_idx, end_idx, section_style, content_styles

            def insert_formatted_content(start_idx, end_idx, content, section_style, content_styles):
                if start_idx is None:
                    return
                
                # Remove old content
                for _ in range(end_idx - start_idx - 1):
                    del doc.paragraphs[start_idx+1]._element
                
                # Split content into paragraphs
                paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
                
                # Insert new content with preserved formatting
                if paragraphs:
                    # First paragraph gets the first content style (or default if none)
                    style = content_styles[0] if content_styles else None
                    p = doc.add_paragraph(paragraphs[0], style=style)
                    p._p.addnext(doc.paragraphs[start_idx]._p)  # Move after section header
                    
                    # Remaining paragraphs get subsequent styles (cycling if needed)
                    for i, para_text in enumerate(paragraphs[1:], 1):
                        style = content_styles[i % len(content_styles)] if content_styles else None
                        p = doc.add_paragraph(para_text, style=style)
                        p._p.addnext(doc.paragraphs[start_idx+1]._p)  # Keep in order

            # Update SUMMARY section
            summary_start, summary_end, summary_style, summary_content_styles = find_section_with_formatting('SUMMARY')
            if summary_start is not None:
                # Preserve the section header
                doc.paragraphs[summary_start].style = summary_style
                insert_formatted_content(summary_start, summary_end, tailored_summary, summary_style, summary_content_styles)

            # Update SKILLS section (try both common names)
            skills_start, skills_end, skills_style, skills_content_styles = find_section_with_formatting('SKILLS')
            if skills_start is None:
                skills_start, skills_end, skills_style, skills_content_styles = find_section_with_formatting('KEY SKILLS')
            
            if skills_start is not None:
                # Preserve the section header
                doc.paragraphs[skills_start].style = skills_style
                # For skills, we want to preserve bullet point formatting
                if not any('•' in line for line in tailored_skills.split('\n')):
                    # If no bullets in input, add them
                    tailored_skills = '\n'.join(f'• {line.strip()}' for line in tailored_skills.split('\n') if line.strip())
                insert_formatted_content(skills_start, skills_end, tailored_skills, skills_style, skills_content_styles)

            # Save the updated document
            updated_docx_path = os.path.join(tmpdir, 'updated.docx')
            doc.save(updated_docx_path)
            
            # Convert to PDF
            pdf_path = convert_to_pdf(updated_docx_path, tmpdir)
            
            if not pdf_path:
                # If PDF conversion failed, return the DOCX instead
                return send_file(
                    updated_docx_path,
                    as_attachment=True,
                    download_name='tailored_resume.docx',
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            
            return send_file(
                pdf_path,
                as_attachment=True,
                download_name='tailored_resume.pdf',
                mimetype='application/pdf'
            )
            
    except Exception as e:
        print(f'Error in /download-docx: {e}')
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# Add a cleanup function to run on server shutdown
@atexit.register
def cleanup_on_exit():
    """Clean up all PDF files in the static/pdfs directory on server shutdown."""
    pdf_dir = os.path.join(app.root_path, 'static', 'pdfs')
    if os.path.exists(pdf_dir):
        try:
            shutil.rmtree(pdf_dir)
            print(f"Cleaned up PDF directory: {pdf_dir}")
        except Exception as e:
            print(f"Error cleaning up PDF directory: {e}")

# --- Main Execution ---
if __name__ == '__main__':
    # Check LaTeX packages before starting the server
    if not check_latex_packages():
        print("Warning: Required LaTeX packages may not be installed.")
        print("Please install the following packages:")
        print("  - texlive-latex-base")
        print("  - texlive-latex-extra")
        print("  - texlive-fonts-recommended")
        print("  - texlive-fonts-extra")
        print("\nOn macOS: brew install basictex")
        print("On Ubuntu/Debian: sudo apt-get install texlive-latex-base texlive-latex-extra texlive-fonts-recommended texlive-fonts-extra")
        print("On Windows: Install MiKTeX or TeX Live")
    
    # Set host='0.0.0.0' to make it accessible on your network (use with caution)
    # Remove debug=True for production environments
    app.run(debug=True, host='127.0.0.1', port=5100)